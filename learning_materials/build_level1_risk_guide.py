from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(__file__).with_name("oTree_Level_1_Risk_Choice_Study_Guide.docx")

NAVY = RGBColor(31, 77, 120)
BLUE = RGBColor(46, 116, 181)
INK = RGBColor(32, 45, 58)
MUTED = RGBColor(95, 105, 115)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CODE_FILL = "F6F8FA"
CAUTION_FILL = "FFF4CE"
WHITE = RGBColor(255, 255, 255)


def set_run_font(run, name="Calibri", size=11, color=INK, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    total = sum(widths_dxa)
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def style_table(table, widths_dxa, header=True):
    table.style = "Table Grid"
    set_table_geometry(table, widths_dxa)
    if header:
        set_repeat_table_header(table.rows[0])
        for cell in table.rows[0].cells:
            shade_cell(cell, LIGHT_BLUE)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, size=10, color=NAVY, bold=True)
    for row_index, row in enumerate(table.rows):
        if row_index and row_index % 2 == 0:
            for cell in row.cells:
                shade_cell(cell, LIGHT_GRAY)
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.1
                for run in paragraph.runs:
                    if not (header and row_index == 0):
                        set_run_font(run, size=10, color=INK)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, value, end])


def add_code(doc, code):
    paragraph = doc.add_paragraph(style="Code Block")
    run = paragraph.add_run(code.strip("\n"))
    set_run_font(run, name="Consolas", size=9, color=INK)
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), CODE_FILL)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), "2E74B5")
    borders.append(left)
    p_pr.append(borders)
    return paragraph


def add_lead(doc, label, text, fill=LIGHT_BLUE):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.12)
    paragraph.paragraph_format.right_indent = Inches(0.12)
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(9)
    paragraph.paragraph_format.line_spacing = 1.15
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "16")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "2E74B5")
    borders.append(left)
    p_pr.append(borders)
    label_run = paragraph.add_run(f"{label}: ")
    set_run_font(label_run, size=10.5, color=NAVY, bold=True)
    text_run = paragraph.add_run(text)
    set_run_font(text_run, size=10.5, color=INK)


def add_body(doc, text, bold_lead=None):
    paragraph = doc.add_paragraph()
    if bold_lead:
        lead = paragraph.add_run(bold_lead)
        set_run_font(lead, bold=True)
    run = paragraph.add_run(text)
    set_run_font(run)
    return paragraph


def configure_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, NAVY, 10, 5),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    code_style = doc.styles.add_style("Code Block", 1)
    code_style.font.name = "Consolas"
    code_style.font.size = Pt(9)
    code_style.paragraph_format.left_indent = Inches(0.12)
    code_style.paragraph_format.right_indent = Inches(0.05)
    code_style.paragraph_format.space_before = Pt(5)
    code_style.paragraph_format.space_after = Pt(7)
    code_style.paragraph_format.line_spacing = 1.0
    code_style.paragraph_format.keep_together = True

    header = section.header.paragraphs[0]
    header.text = "oTree Learning Series | Level 1"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in header.runs:
        set_run_font(run, size=9, color=MUTED, bold=True)
    add_page_number(section.footer.paragraphs[0])
    return doc


def add_cover(doc):
    doc.add_paragraph().paragraph_format.space_after = Pt(42)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = kicker.add_run("OTREE LEARNING SERIES")
    set_run_font(run, size=11, color=BLUE, bold=True)
    kicker.paragraph_format.space_after = Pt(18)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Level 1: Individual Choice Under Risk")
    set_run_font(title_run, size=28, color=NAVY, bold=True)
    title.paragraph_format.space_after = Pt(10)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(
        "A code-by-code guide to Player fields, payoff resolution, Pages, and HTML templates"
    )
    set_run_font(subtitle_run, size=14, color=MUTED)
    subtitle.paragraph_format.space_after = Pt(34)

    rule = doc.add_paragraph()
    p_pr = rule._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:color"), "2E74B5")
    borders.append(bottom)
    p_pr.append(borders)
    rule.paragraph_format.space_after = Pt(30)

    details = doc.add_table(rows=1, cols=2)
    details.rows[0].cells[0].text = "Document detail"
    details.rows[0].cells[1].text = "Value"
    rows = [
        ("Project", "otree_marielle"),
        ("App", "learning_1_risk"),
        ("Framework", "oTree 6.0.15"),
        ("Purpose", "Study reference and future revision notes"),
    ]
    for values in rows:
        row = details.add_row()
        row.cells[0].text = values[0]
        row.cells[1].text = values[1]
    style_table(details, [2700, 6660], header=True)
    for row in details.rows[1:]:
        shade_cell(row.cells[0], LIGHT_BLUE)
        for run in row.cells[0].paragraphs[0].runs:
            set_run_font(run, size=10, color=NAVY, bold=True)
    doc.add_page_break()


def build_document():
    doc = configure_document()
    add_cover(doc)

    doc.add_heading("1. Experiment overview", level=1)
    add_body(
        doc,
        "This experiment asks one participant to choose between a certain payment and a risky lottery. "
        "The participant first reads the instructions, then submits a choice. oTree stores the choice, "
        "runs the server-side resolver, stores the random draw and payoff, and renders the result page."
    )
    add_lead(
        doc,
        "Central lesson",
        "Models define stored data, Pages control timing and processing, ordinary Python functions implement "
        "economic logic, and HTML templates control what participants see.",
    )

    doc.add_heading("Execution flow", level=2)
    flow = doc.add_table(rows=1, cols=3)
    flow.rows[0].cells[0].text = "Stage"
    flow.rows[0].cells[1].text = "Python/oTree action"
    flow.rows[0].cells[2].text = "Participant-facing result"
    flow_rows = [
        ("1", "Introduction Page opens", "Instructions and Next button"),
        ("2", "Decision uses Player.choice", "Safe/risky radio buttons"),
        ("3", "Valid form is saved", "The selected value reaches the server"),
        ("4", "before_next_page calls resolve_choice", "No separate screen; calculation happens server-side"),
        ("5", "Results.vars_for_template prepares display values", "Readable choice and total are available"),
        ("6", "Decision.html renders stored and temporary values", "Draw, payoff, and total are shown"),
    ]
    for values in flow_rows:
        cells = flow.add_row().cells
        for cell, value in zip(cells, values):
            cell.text = value
    style_table(flow, [900, 3960, 4500])

    doc.add_heading("2. Imports and app description", level=1)
    add_code(
        doc,
        """import random

from otree.api import (
    BaseConstants, BaseGroup, BasePlayer, BaseSubsession,
    Currency, Page, models, widgets,
)""",
    )
    import_table = doc.add_table(rows=1, cols=2)
    import_table.rows[0].cells[0].text = "Imported name"
    import_table.rows[0].cells[1].text = "Purpose in this app"
    import_rows = [
        ("random", "Python's random-number module; provides randint(1, 100)."),
        ("BaseConstants", "Parent class for fixed experiment parameters."),
        ("BaseSubsession", "Parent class for one round of the app."),
        ("BaseGroup", "Parent class for group-level records."),
        ("BasePlayer", "Parent class for one participant's app-and-round record."),
        ("Currency", "Creates values that oTree treats and formats as money."),
        ("Page", "Parent class for participant screens."),
        ("models", "Provides database field types."),
        ("widgets", "Controls form appearance, including radio buttons."),
    ]
    for values in import_rows:
        cells = import_table.add_row().cells
        cells[0].text, cells[1].text = values
    style_table(import_table, [2400, 6960])
    add_body(
        doc,
        "The doc string describes the app for researchers and administrators. It documents the design but "
        "does not change participant flow, stored data, or payoffs."
    )

    doc.add_heading("3. Constants and oTree hierarchy", level=1)
    add_code(
        doc,
        """class C(BaseConstants):
    NAME_IN_URL = "learning-risk"
    PLAYERS_PER_GROUP = None
    NUM_ROUNDS = 1
    ENDOWMENT = Currency(5)
    SAFE_AMOUNT = Currency(4)
    HIGH_PRIZE = Currency(10)
    LOW_PRIZE = Currency(0)
    HIGH_PRIZE_PROBABILITY = 60""",
    )
    constants = doc.add_table(rows=1, cols=3)
    constants.rows[0].cells[0].text = "Constant"
    constants.rows[0].cells[1].text = "Current value"
    constants.rows[0].cells[2].text = "Effect"
    constant_rows = [
        ("NAME_IN_URL", "learning-risk", "Sets the app-specific URL fragment."),
        ("PLAYERS_PER_GROUP", "None", "Creates an individual rather than fixed-size multiplayer task."),
        ("NUM_ROUNDS", "1", "Runs the complete page sequence once."),
        ("ENDOWMENT", "GBP 5", "Displayed starting endowment."),
        ("SAFE_AMOUNT", "GBP 4", "Decision payoff for the safe option."),
        ("HIGH_PRIZE", "GBP 10", "High lottery outcome."),
        ("LOW_PRIZE", "GBP 0", "Low lottery outcome."),
        ("HIGH_PRIZE_PROBABILITY", "60", "Winning draws are 1 through 60."),
    ]
    for values in constant_rows:
        cells = constants.add_row().cells
        for cell, value in zip(cells, values):
            cell.text = value
    style_table(constants, [2700, 1500, 5160])
    add_body(
        doc,
        "Subsession represents one app round and Group represents a set of players in that round. Both are "
        "declared with pass because this individual experiment needs no custom round-level or group-level fields."
    )
    add_code(
        doc,
        """class Subsession(BaseSubsession):
    pass

class Group(BaseGroup):
    pass""",
    )

    doc.add_page_break()
    doc.add_heading("4. Deep dive: the Player class", level=1)
    add_lead(
        doc,
        "Definition",
        "A Player is one participant's database record for one app and one round. Participant persists across "
        "apps; Player is app-and-round specific.",
    )
    add_code(
        doc,
        """class Player(BasePlayer):
    choice = models.StringField(
        label="Which option do you choose?",
        choices=[
            ["safe", "Option A: receive GBP 4 for certain"],
            ["risky", "Option B: 60% chance of GBP 10, otherwise GBP 0"],
        ],
        widget=widgets.RadioSelect,
    )
    random_draw = models.IntegerField()
    lottery_won = models.BooleanField()""",
    )

    doc.add_heading("4.1 Player.choice", level=2)
    add_body(
        doc,
        "choice is a StringField, so oTree creates a database column that stores text. It becomes a participant "
        "input only because Decision lists it in form_fields. The field is required by default."
    )
    mapping = doc.add_table(rows=1, cols=2)
    mapping.rows[0].cells[0].text = "Participant sees"
    mapping.rows[0].cells[1].text = "Database stores"
    for displayed, stored in [
        ("Option A: receive GBP 4 for certain", '"safe"'),
        ("Option B: 60% chance of GBP 10, otherwise GBP 0", '"risky"'),
    ]:
        cells = mapping.add_row().cells
        cells[0].text, cells[1].text = displayed, stored
    style_table(mapping, [6500, 2860])
    add_body(
        doc,
        "label supplies the question text. choices contains pairs of stored value and displayed label. "
        "RadioSelect renders mutually exclusive radio buttons. The resolver must compare player.choice with "
        "the stored value, such as \"safe\", not the long visible label."
    )

    doc.add_heading("4.2 Player.random_draw", level=2)
    add_body(
        doc,
        "random_draw is an IntegerField. The participant never enters it because it is absent from form_fields. "
        "The server assigns a number from 1 through 100. Because it is a model field, the value is retained in "
        "the database and appears in oTree exports."
    )

    doc.add_heading("4.3 Player.lottery_won", level=2)
    add_body(
        doc,
        "lottery_won is a BooleanField. For a risky choice it records whether the draw won. For a safe choice the "
        "current code stores False, meaning the lottery was not won or used. A production design could add a "
        "separate lottery_played field if it must distinguish not played from played-and-lost."
    )

    doc.add_heading("4.4 Built-in Player.payoff", level=2)
    add_body(
        doc,
        "payoff is inherited from BasePlayer, so it is not declared in the custom Player class. oTree aggregates "
        "round payoffs into participant.payoff. Assigning a custom field does not affect payment; assigning "
        "player.payoff does."
    )
    add_code(doc, "player.payoff = C.SAFE_AMOUNT")

    doc.add_heading("5. Deep dive: resolve_choice", level=1)
    add_code(
        doc,
        """def resolve_choice(player: Player):
    \"\"\"Make one server-side random draw and set this round's payoff.\"\"\"
    player.random_draw = random.randint(1, 100)
    if player.choice == "safe":
        player.lottery_won = False
        player.payoff = C.SAFE_AMOUNT
    else:
        player.lottery_won = player.random_draw <= C.HIGH_PRIZE_PROBABILITY
        player.payoff = C.HIGH_PRIZE if player.lottery_won else C.LOW_PRIZE""",
    )
    add_body(
        doc,
        "resolve_choice is an ordinary Python function, not an automatic oTree hook. It runs only because "
        "Decision.before_next_page explicitly calls it. The Player type annotation helps PyCharm and readers; "
        "it does not create the Player record."
    )

    resolver = doc.add_table(rows=1, cols=3)
    resolver.rows[0].cells[0].text = "Statement"
    resolver.rows[0].cells[1].text = "What it does"
    resolver.rows[0].cells[2].text = "Stored consequence"
    resolver_rows = [
        ("random.randint(1, 100)", "Generates an inclusive whole-number draw on the server.", "random_draw is populated."),
        ('player.choice == "safe"', "Checks the short value saved by the form.", "Selects safe or risky branch."),
        ("draw <= 60", "Maps exactly 60 of 100 possible draws to a win.", "lottery_won becomes True/False."),
        ("A if condition else B", "Python conditional expression chooses one of two prizes.", "payoff becomes GBP 10 or GBP 0."),
    ]
    for values in resolver_rows:
        cells = resolver.add_row().cells
        for cell, value in zip(cells, values):
            cell.text = value
    style_table(resolver, [2800, 3500, 3060])

    doc.add_heading("5.1 Safe branch", level=2)
    add_code(
        doc,
        """if player.choice == "safe":
    player.lottery_won = False
    player.payoff = C.SAFE_AMOUNT""",
    )
    add_body(
        doc,
        "A safe-choice record may contain choice=\"safe\", random_draw=37, lottery_won=False, and payoff=GBP 4. "
        "The draw is still made and stored, but it has no payment consequence."
    )

    doc.add_heading("5.2 Risky branch", level=2)
    add_code(
        doc,
        """player.lottery_won = player.random_draw <= C.HIGH_PRIZE_PROBABILITY
player.payoff = C.HIGH_PRIZE if player.lottery_won else C.LOW_PRIZE""",
    )
    add_body(
        doc,
        "Draws 1-60 win and draws 61-100 lose. A draw of 23 produces lottery_won=True and payoff=GBP 10; "
        "a draw of 84 produces lottery_won=False and payoff=GBP 0."
    )
    add_lead(
        doc,
        "Design benefit",
        "Keeping economic logic in a named function makes it easier to read, test, reuse, and separate from page flow.",
    )

    doc.add_page_break()
    doc.add_heading("6. Pages: commands and timing", level=1)
    add_body(
        doc,
        "Each Page class is matched to an HTML file with the same name. Page classes control data collection and "
        "processing; templates control presentation."
    )

    doc.add_heading("6.1 Introduction", level=2)
    add_code(
        doc,
        """class Introduction(Page):
    pass""",
    )
    add_body(
        doc,
        "The page has no form fields or special hooks, so pass is sufficient. oTree automatically looks for "
        "Introduction.html. Clicking Next submits an empty page and advances."
    )

    doc.add_heading("6.2 Decision", level=2)
    add_code(
        doc,
        """class Decision(Page):
    form_model = "player"
    form_fields = ["choice"]

    @staticmethod
    def before_next_page(player: Player, timeout_happened):
        resolve_choice(player)""",
    )
    page_commands = doc.add_table(rows=1, cols=2)
    page_commands.rows[0].cells[0].text = "Command"
    page_commands.rows[0].cells[1].text = "Meaning"
    command_rows = [
        ('form_model = "player"', "Save submitted values on the current Player record."),
        ('form_fields = ["choice"]', "Build and validate a form for Player.choice."),
        ("@staticmethod", "The hook needs no Decision object; oTree supplies Player directly."),
        ("before_next_page", "Runs after successful validation/save and before Results opens."),
        ("timeout_happened", "Indicates timer submission; unused because this page has no timeout."),
        ("resolve_choice(player)", "Calls the ordinary Python resolver with the current Player record."),
    ]
    for values in command_rows:
        cells = page_commands.add_row().cells
        cells[0].text, cells[1].text = values
    style_table(page_commands, [3300, 6060])
    add_lead(
        doc,
        "Critical ordering",
        "oTree validates and saves choice before before_next_page runs. Therefore resolve_choice can safely read "
        "player.choice.",
    )

    doc.add_heading("6.3 Results", level=2)
    add_code(
        doc,
        """class Results(Page):
    @staticmethod
    def vars_for_template(player: Player):
        return dict(
            chosen_option="safe amount" if player.choice == "safe" else "lottery",
            total_payment=C.ENDOWMENT + player.payoff,
        )""",
    )
    add_body(
        doc,
        "vars_for_template returns display-only values. chosen_option converts an internal code into readable text. "
        "total_payment adds the displayed endowment to the built-in decision payoff. These names are available to "
        "Decision.html but are not new database columns. The method runs again if the page is refreshed."
    )

    doc.add_heading("6.4 page_sequence", level=2)
    add_code(doc, "page_sequence = [Introduction, Decision, Results]")
    add_body(
        doc,
        "page_sequence is the exact participant flow. The entries are Page class references, not strings. A Page "
        "that is not listed is not shown. With NUM_ROUNDS=1, the sequence runs once."
    )

    doc.add_heading("7. HTML templates and their Python connections", level=1)
    html_commands = doc.add_table(rows=1, cols=2)
    html_commands.rows[0].cells[0].text = "Template command"
    html_commands.rows[0].cells[1].text = "Function"
    html_rows = [
        ("{{ block title }}", "Fills the title area of oTree's standard page."),
        ("{{ block content }}", "Fills the main page/form area."),
        ("{{ C.NAME }}", "Reads a fixed parameter from the C class."),
        ("{{ formfields }}", "Renders every field listed in the Page's form_fields."),
        ("{{ next_button }}", "Renders the oTree form submit button."),
        ("{{ player.field }}", "Reads a stored field from the current Player record."),
        ("{{ chosen_option }}", "Reads an extra value returned by vars_for_template."),
    ]
    for values in html_rows:
        cells = html_commands.add_row().cells
        cells[0].text, cells[1].text = values
    style_table(html_commands, [3200, 6160])

    doc.add_heading("7.1 Introduction.html", level=2)
    add_code(
        doc,
        """{{ block title }}Level 1: choice under risk{{ endblock }}
{{ block content }}
<p>You receive a participation endowment of {{ C.ENDOWMENT }}.</p>
{{ next_button }}
{{ endblock }}""",
    )
    add_body(
        doc,
        "C.ENDOWMENT is automatically available in the template and is formatted as project currency. Bootstrap "
        "classes on the div provide card styling without changing experimental logic."
    )

    doc.add_heading("7.2 Decision.html", level=2)
    add_code(
        doc,
        """<p>The risky option pays {{ C.HIGH_PRIZE }} with probability
   {{ C.HIGH_PRIZE_PROBABILITY }}%, and {{ C.LOW_PRIZE }} otherwise.</p>

{{ formfields }}
{{ next_button }}""",
    )
    add_body(
        doc,
        "formfields connects to Decision.form_fields and Player.choice. oTree uses the field label, choices, and "
        "RadioSelect widget to generate the question, radio inputs, submitted values, and validation errors."
    )

    doc.add_heading("7.3 Decision.html", level=2)
    add_code(
        doc,
        """<tr><th>Your choice</th><td>{{ chosen_option }}</td></tr>
<tr><th>Random draw</th><td>{{ player.random_draw }}</td></tr>
<tr><th>Decision payoff</th><td>{{ player.payoff }}</td></tr>
<tr><th>Endowment + decision payoff</th><td>{{ total_payment }}</td></tr>""",
    )
    results_map = doc.add_table(rows=1, cols=3)
    results_map.rows[0].cells[0].text = "HTML expression"
    results_map.rows[0].cells[1].text = "Python source"
    results_map.rows[0].cells[2].text = "Database field?"
    result_rows = [
        ("chosen_option", "Results.vars_for_template", "No"),
        ("player.random_draw", "Player.random_draw", "Yes"),
        ("player.payoff", "Inherited BasePlayer.payoff", "Yes"),
        ("total_payment", "Results.vars_for_template", "No"),
        ("C.ENDOWMENT", "C constants class", "Fixed parameter"),
    ]
    for values in result_rows:
        cells = results_map.add_row().cells
        for cell, value in zip(cells, values):
            cell.text = value
    style_table(results_map, [3200, 3900, 2260])

    doc.add_page_break()
    doc.add_heading("8. Complete Python-to-HTML chain", level=1)
    chain = doc.add_table(rows=1, cols=2)
    chain.rows[0].cells[0].text = "Layer"
    chain.rows[0].cells[1].text = "Responsibility"
    chain_rows = [
        ("Player model", "Declares choice, random_draw, and lottery_won; inherits payoff."),
        ("Decision Page", "Chooses Player as form model and exposes choice as an input."),
        ("Decision.html", "Renders choice via formfields and submits it via next_button."),
        ("oTree form processing", "Validates the allowed value and saves player.choice."),
        ("before_next_page", "Calls resolve_choice after the valid save."),
        ("resolve_choice", "Stores draw, win status, and official round payoff."),
        ("Results Page", "Creates readable display-only variables."),
        ("Decision.html", "Displays stored Player fields and temporary template values."),
    ]
    for values in chain_rows:
        cells = chain.add_row().cells
        cells[0].text, cells[1].text = values
    style_table(chain, [2600, 6760])

    doc.add_heading("9. Important payment distinction", level=1)
    add_lead(
        doc,
        "Warning",
        "The current Results page displays ENDOWMENT + player.payoff, but the built-in oTree payoff contains only "
        "the decision payoff. A displayed total of GBP 9 can coexist with an official oTree payoff of GBP 4.",
        fill=CAUTION_FILL,
    )
    add_body(
        doc,
        "If the endowment should be paid as part of this app, include it when assigning player.payoff. If it is a "
        "show-up fee, configure it as participation_fee in the session config. Payment rules must match participant "
        "instructions and exported payment data before a real study is launched."
    )
    add_code(doc, "player.payoff = C.ENDOWMENT + C.SAFE_AMOUNT")

    doc.add_heading("10. Automated bot", level=1)
    add_code(
        doc,
        """class PlayerBot(Bot):
    def play_round(self):
        yield Introduction
        yield Decision, dict(choice="safe")
        assert self.player.payoff == C.SAFE_AMOUNT
        yield Submission(Results, check_html=False)""",
    )
    add_body(
        doc,
        "yield Introduction visits the page. The Decision yield submits the stored value \"safe\". Normal form "
        "processing and before_next_page still run. The assertion verifies the safe payoff. Submission visits the "
        "Results page while disabling detailed HTML checking for that page."
    )

    doc.add_heading("11. Study checklist", level=1)
    checklist = doc.add_table(rows=1, cols=2)
    checklist.rows[0].cells[0].text = "Check"
    checklist.rows[0].cells[1].text = "Question to answer"
    checks = [
        ("Stored values", "Can you identify which values are database fields and which are display-only?"),
        ("Form mapping", "Does every form_fields name exist on the stated form_model?"),
        ("Timing", "Does the calculation run only after the required values have been saved?"),
        ("Randomization", "Is the random draw performed on the server and mapped correctly to probability?"),
        ("Payoff", "Does built-in player.payoff equal the amount the participant is actually promised?"),
        ("Template", "Can each {{ variable }} be traced to C, Player, or vars_for_template?"),
        ("Data", "Will every variable needed for analysis appear in the export?"),
        ("Test", "Does a bot verify at least one safe and one risky boundary case?"),
    ]
    for values in checks:
        cells = checklist.add_row().cells
        cells[0].text, cells[1].text = values
    style_table(checklist, [2200, 7160])

    doc.add_heading("12. Source files", level=1)
    add_body(doc, "Python app: learning_1_risk/__init__.py", bold_lead="Primary model and flow - ")
    add_body(doc, "learning_1_risk/Introduction.html", bold_lead="Instructions template - ")
    add_body(doc, "learning_1_risk/Decision.html", bold_lead="Decision template - ")
    add_body(doc, "learning_1_risk/Decision.html", bold_lead="Results template - ")
    add_body(doc, "learning_1_risk/tests.py", bold_lead="Automated participant - ")

    doc.core_properties.title = "oTree Level 1: Individual Choice Under Risk"
    doc.core_properties.subject = "oTree code and HTML study guide"
    doc.core_properties.author = "oTree Learning Series"
    doc.core_properties.keywords = "oTree, behavioral economics, risk, Player, Page, HTML"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
